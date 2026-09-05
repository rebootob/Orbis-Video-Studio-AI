import io
import time
import uuid
from typing import Dict, Any, List
import docx
from app.core.config import settings
from app.services.document_extraction.base import DocumentExtractor, ExtractionResult


class DocxDocumentExtractor(DocumentExtractor):
    """Fast native Word document (.docx) extractor using python-docx."""

    def supports(self, document_type: str) -> bool:
        return document_type == "docx"

    def extract(self, asset_id: uuid.UUID, content: bytes, metadata: Dict[str, Any]) -> ExtractionResult:
        start_time = time.perf_counter()
        warnings: List[str] = []

        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception as e:
            duration_ms = round((time.perf_counter() - start_time) * 1000, 2)
            return ExtractionResult(
                asset_id=asset_id,
                document_type="docx",
                extracted_text="",
                character_count=0,
                segment_count=0,
                extraction_status="FAILED",
                extraction_method="python-docx",
                extraction_duration_ms=duration_ms,
                error_message=f"Failed to parse DOCX document: {str(e)}",
            )

        paragraphs_text: List[str] = []
        segments: List[Dict[str, Any]] = []
        seg_idx = 1
        total_chars = 0

        # Extract document paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                text_clean = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
                paragraphs_text.append(text_clean)
                segments.append({
                    "index": seg_idx,
                    "type": "paragraph",
                    "text": text_clean,
                })
                seg_idx += 1
                total_chars += len(text_clean)

                if total_chars >= settings.MAX_EXTRACTED_CHARACTERS:
                    warnings.append(f"Text extraction limit ({settings.MAX_EXTRACTED_CHARACTERS} chars) reached.")
                    break

        # Extract document table cell text if character limit not exceeded
        if total_chars < settings.MAX_EXTRACTED_CHARACTERS:
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs_text.append(row_text)
                        segments.append({
                            "index": seg_idx,
                            "type": "table_row",
                            "text": row_text,
                        })
                        seg_idx += 1
                        total_chars += len(row_text)
                        if total_chars >= settings.MAX_EXTRACTED_CHARACTERS:
                            break

        combined_text = "\n\n".join(paragraphs_text)
        duration_ms = round((time.perf_counter() - start_time) * 1000, 2)
        status = "SUCCESS" if combined_text.strip() else "NO_TEXT_LAYER"

        return ExtractionResult(
            asset_id=asset_id,
            document_type="docx",
            extracted_text=combined_text,
            character_count=len(combined_text),
            segment_count=len(segments),
            segments=segments,
            extraction_status=status,
            extraction_method="python-docx",
            extraction_duration_ms=duration_ms,
            warnings=warnings,
        )
